"""Shared fixtures for document pipeline tests."""

from __future__ import annotations

from pathlib import Path

import pytest


FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session")
def fixtures_dir() -> Path:
    """Return the fixtures directory, creating sample files if needed."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    _ensure_markdown_fixture()
    _ensure_txt_fixtures()
    _ensure_docx_fixture()
    _ensure_pdf_fixture()
    return FIXTURES_DIR


@pytest.fixture
def markdown_file(fixtures_dir: Path) -> Path:
    return fixtures_dir / "sample.md"


@pytest.fixture
def txt_utf8_file(fixtures_dir: Path) -> Path:
    return fixtures_dir / "sample_utf8.txt"


@pytest.fixture
def txt_gbk_file(fixtures_dir: Path) -> Path:
    return fixtures_dir / "sample_gbk.txt"


@pytest.fixture
def docx_file(fixtures_dir: Path) -> Path:
    return fixtures_dir / "sample.docx"


@pytest.fixture
def pdf_file(fixtures_dir: Path) -> Path:
    return fixtures_dir / "sample.pdf"


def _ensure_markdown_fixture() -> None:
    path = FIXTURES_DIR / "sample.md"
    if path.exists():
        return
    path.write_text(
        """---
title: Kubernetes部署规范
author: SRE团队
tags:
  - k8s
  - docker
---

# Kubernetes部署规范

本文档描述 Kubernetes 集群部署标准。

## 前置条件

- 节点已安装 containerd
""",
        encoding="utf-8",
    )


def _ensure_txt_fixtures() -> None:
    utf8_path = FIXTURES_DIR / "sample_utf8.txt"
    if not utf8_path.exists():
        utf8_path.write_text("UTF-8 文本内容：Pod OOM 排查手册\n第二行", encoding="utf-8")

    gbk_path = FIXTURES_DIR / "sample_gbk.txt"
    if not gbk_path.exists():
        gbk_path.write_bytes("GBK编码内容：服务器重启流程\n第二行".encode("gbk"))


def _ensure_docx_fixture() -> None:
    path = FIXTURES_DIR / "sample.docx"
    if path.exists():
        return
    from docx import Document

    doc = Document()
    doc.add_heading("Kubernetes故障排查", level=1)
    doc.add_heading("Pod OOM", level=2)
    doc.add_paragraph("当 Pod 因内存不足被 Kill 时，按以下步骤排查。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "CPU"
    table.rows[1].cells[1].text = "80%"
    doc.core_properties.author = "SRE Team"
    doc.save(str(path))


def _ensure_pdf_fixture() -> None:
    path = FIXTURES_DIR / "sample.pdf"
    if path.exists():
        return
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Enterprise Knowledge Sample PDF")
    page.insert_text((72, 100), "Page 1 content for DevOps SOP.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page 2 content.")
    doc.set_metadata(
        {
            "title": "Sample PDF Doc",
            "author": "Knowledge Copilot",
            "creator": "PyMuPDF Test",
        }
    )
    doc.save(str(path))
    doc.close()
