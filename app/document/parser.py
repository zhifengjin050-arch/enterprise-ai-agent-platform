"""
Document parser module.

Parses enterprise knowledge documents (PDF, Word, Markdown, TXT)
into a unified ParsedDocument structure for downstream workflow processing.
"""

from __future__ import annotations

import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Union

from pydantic import BaseModel, Field


class UnsupportedFormatError(ValueError):
    """Raised when a document format is not supported."""

    def __init__(self, format_or_path: str, supported: Optional[List[str]] = None) -> None:
        self.format_or_path = format_or_path
        self.supported = supported or [".pdf", ".docx", ".md", ".markdown", ".txt"]
        super().__init__(
            f"Unsupported file format: {format_or_path}. "
            f"Supported formats: {', '.join(self.supported)}"
        )


class DocumentParseError(RuntimeError):
    """Raised when document parsing fails for a supported format."""

    def __init__(self, path: str, reason: str) -> None:
        self.path = path
        self.reason = reason
        super().__init__(f"Failed to parse document '{path}': {reason}")


class ParsedDocument(BaseModel):
    """Unified document representation after parsing.

    All document formats are normalized into this structure before
    conversion to Markdown and submission to the knowledge workflow.
    """

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str = ""
    content: str = ""
    format: str = "text"
    metadata: Dict[str, Any] = Field(default_factory=dict)
    source: str = ""
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


def _table_to_markdown(rows: List[List[str]]) -> str:
    """Convert a 2D table into a Markdown table string.

    Args:
        rows: List of rows, each row is a list of cell strings.

    Returns:
        Markdown table text. Empty string if rows is empty.
    """
    if not rows:
        return ""

    normalized: List[List[str]] = [
        [str(cell).replace("\n", " ").strip() for cell in row] for row in rows
    ]
    col_count = max(len(row) for row in normalized)
    for row in normalized:
        while len(row) < col_count:
            row.append("")

    header = normalized[0]
    body = normalized[1:] if len(normalized) > 1 else []

    lines = [
        "|" + "|".join(header) + "|",
        "|" + "|".join(["-"] * col_count) + "|",
    ]
    for row in body:
        lines.append("|" + "|".join(row) + "|")
    return "\n".join(lines)


def parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF file with page structure, tables, and image placeholders.

    Uses PyMuPDF (fitz). Tables are converted to Markdown. Images are
    recorded as metadata only (no OCR).

    Args:
        file_path: Path to the PDF file.

    Returns:
        ParsedDocument with extracted content and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        DocumentParseError: If PDF parsing fails.
        ImportError: If PyMuPDF is not installed.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required for PDF parsing: pip install PyMuPDF"
        ) from exc

    try:
        doc = fitz.open(str(file_path))
    except Exception as exc:
        raise DocumentParseError(str(file_path), str(exc)) from exc

    content_parts: List[str] = []
    images: List[Dict[str, Any]] = []
    table_count = 0

    try:
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_lines: List[str] = []

            # Page marker for structure preservation
            page_lines.append(f"<!-- page:{page_num + 1} -->")

            # Extract tables first (if supported by installed PyMuPDF)
            table_rects: List[Any] = []
            if hasattr(page, "find_tables"):
                try:
                    finder = page.find_tables()
                    tables = getattr(finder, "tables", finder) if finder else []
                    for table in tables:
                        try:
                            extracted = table.extract()
                            if extracted:
                                md_table = _table_to_markdown(extracted)
                                if md_table:
                                    page_lines.append(md_table)
                                    table_count += 1
                            if hasattr(table, "bbox"):
                                table_rects.append(fitz.Rect(table.bbox))
                        except Exception:
                            # Skip individual table extraction failures
                            continue
                except Exception:
                    # find_tables not usable on this page/version
                    pass

            # Extract text blocks outside table regions when possible
            text = page.get_text("text")
            if text and text.strip():
                page_lines.append(text.strip())

            # Image metadata only (no OCR)
            for img_index, img in enumerate(page.get_images(full=True)):
                images.append({"page": page_num + 1, "index": img_index})

            content_parts.append("\n\n".join(page_lines))

        pdf_meta = doc.metadata or {}
        metadata: Dict[str, Any] = {
            "pages": doc.page_count,
            "page_count": doc.page_count,
            "author": pdf_meta.get("author") or "",
            "creator": pdf_meta.get("creator") or "",
            "subject": pdf_meta.get("subject") or "",
            "producer": pdf_meta.get("producer") or "",
            "tables": table_count,
            "images": images,
        }
        title = (pdf_meta.get("title") or "").strip() or file_path.stem
        content = "\n\n".join(content_parts)
    finally:
        doc.close()

    return ParsedDocument(
        title=title,
        content=content,
        format="pdf",
        metadata=metadata,
        source=str(file_path),
    )


def _docx_heading_level(style_name: str) -> Optional[int]:
    """Extract heading level from a Word paragraph style name.

    Args:
        style_name: Style name such as 'Heading 1' or '标题 2'.

    Returns:
        Heading level 1-9, or None if not a heading style.
    """
    name = (style_name or "").strip().lower()
    for prefix in ("heading ", "标题 "):
        if name.startswith(prefix):
            suffix = name[len(prefix) :].strip()
            if suffix.isdigit():
                level = int(suffix)
                if 1 <= level <= 9:
                    return level
    if name in {"title", "标题"}:
        return 1
    return None


def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a Word document preserving heading hierarchy and tables.

    Uses python-docx. Headings become Markdown heading markers.
    Tables become Markdown tables.

    Args:
        file_path: Path to the .docx file.

    Returns:
        ParsedDocument with structured Markdown-oriented content.

    Raises:
        FileNotFoundError: If the file does not exist.
        DocumentParseError: If Word parsing fails.
        ImportError: If python-docx is not installed.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Word file not found: {file_path}")

    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word parsing: pip install python-docx"
        ) from exc

    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        raise DocumentParseError(str(file_path), str(exc)) from exc

    content_parts: List[str] = []
    paragraph_count = 0
    table_count = 0
    title = file_path.stem

    def iter_block_items(parent: Any) -> Any:
        """Yield paragraphs and tables in document order."""
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            paragraph_count += 1
            style_name = block.style.name if block.style is not None else ""
            level = _docx_heading_level(style_name)
            if level is not None:
                content_parts.append(f"{'#' * level} {text}")
                if level == 1 and title == file_path.stem:
                    title = text
            else:
                content_parts.append(text)
        else:
            # Table
            rows: List[List[str]] = []
            for row in block.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            md_table = _table_to_markdown(rows)
            if md_table:
                content_parts.append(md_table)
                table_count += 1

    core = doc.core_properties
    author = getattr(core, "author", None) or ""
    if getattr(core, "title", None):
        title = core.title or title

    metadata: Dict[str, Any] = {
        "author": author,
        "paragraph_count": paragraph_count,
        "tables": table_count,
    }

    return ParsedDocument(
        title=title,
        content="\n\n".join(content_parts),
        format="docx",
        metadata=metadata,
        source=str(file_path),
    )


def parse_markdown(file_path: Path) -> ParsedDocument:
    """Parse a Markdown file and extract YAML frontmatter.

    Uses python-frontmatter. Frontmatter fields become metadata;
    body becomes content.

    Args:
        file_path: Path to the .md / .markdown file.

    Returns:
        ParsedDocument with frontmatter metadata and body content.

    Raises:
        FileNotFoundError: If the file does not exist.
        DocumentParseError: If Markdown parsing fails.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {file_path}")

    try:
        import frontmatter
    except ImportError as exc:
        raise ImportError(
            "python-frontmatter is required for Markdown parsing: "
            "pip install python-frontmatter"
        ) from exc

    try:
        with file_path.open("r", encoding="utf-8") as handle:
            post = frontmatter.load(handle)
    except UnicodeDecodeError:
        try:
            with file_path.open("r", encoding="gbk") as handle:
                post = frontmatter.load(handle)
        except Exception as exc:
            raise DocumentParseError(str(file_path), str(exc)) from exc
    except Exception as exc:
        raise DocumentParseError(str(file_path), str(exc)) from exc

    raw_meta = dict(post.metadata) if post.metadata else {}
    tags = raw_meta.get("tags", [])
    if isinstance(tags, str):
        tags = [t.strip() for t in tags.split(",") if t.strip()]
    elif not isinstance(tags, list):
        tags = list(tags) if tags else []

    metadata: Dict[str, Any] = {
        "title": raw_meta.get("title", ""),
        "author": raw_meta.get("author", ""),
        "tags": tags,
    }
    # Preserve extra frontmatter keys
    for key, value in raw_meta.items():
        if key not in metadata:
            metadata[key] = value

    title = str(raw_meta.get("title") or file_path.stem)

    return ParsedDocument(
        title=title,
        content=post.content.strip(),
        format="markdown",
        metadata=metadata,
        source=str(file_path),
    )


def parse_txt(file_path: Path) -> ParsedDocument:
    """Parse a plain text file with UTF-8 / GBK auto-detection.

    Tries UTF-8 first, then falls back to GBK.

    Args:
        file_path: Path to the .txt file.

    Returns:
        ParsedDocument with decoded content.

    Raises:
        FileNotFoundError: If the file does not exist.
        DocumentParseError: If decoding fails for both encodings.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Text file not found: {file_path}")

    content: Optional[str] = None
    encoding_used = "utf-8"
    last_error: Optional[Exception] = None

    for encoding in ("utf-8", "gbk"):
        try:
            content = file_path.read_text(encoding=encoding)
            encoding_used = encoding
            break
        except UnicodeDecodeError as exc:
            last_error = exc
            continue

    if content is None:
        raise DocumentParseError(
            str(file_path),
            f"Unable to decode as utf-8 or gbk: {last_error}",
        )

    return ParsedDocument(
        title=file_path.stem,
        content=content,
        format="text",
        metadata={"encoding": encoding_used},
        source=str(file_path),
    )


# Backward-compatible alias
parse_text = parse_txt


class DocumentParser:
    """Unified document parser for enterprise knowledge ingestion.

    Detects file format by extension and delegates to the appropriate
    format-specific parser.
    """

    SUPPORTED_EXTENSIONS: Dict[str, Callable[[Path], ParsedDocument]] = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".md": parse_markdown,
        ".markdown": parse_markdown,
        ".txt": parse_txt,
    }

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Parse a document file into a ParsedDocument.

        Args:
            file_path: Path to the document file.

        Returns:
            ParsedDocument with unified structure.

        Raises:
            FileNotFoundError: If the file does not exist.
            UnsupportedFormatError: If the extension is not supported.
            DocumentParseError: If parsing fails.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        extension = path.suffix.lower()
        parser_fn = self.SUPPORTED_EXTENSIONS.get(extension)
        if parser_fn is None:
            raise UnsupportedFormatError(extension, list(self.SUPPORTED_EXTENSIONS.keys()))

        return parser_fn(path)


def parse_file(file_path: Union[str, Path]) -> ParsedDocument:
    """Auto-detect format and parse a document file.

    Convenience wrapper around DocumentParser.parse().

    Args:
        file_path: Path to the document file.

    Returns:
        ParsedDocument with extracted content.
    """
    return DocumentParser().parse(file_path)
